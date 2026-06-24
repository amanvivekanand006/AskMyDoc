from . import *
from . database import *
from . vector import *

load_dotenv()
#------------------secrect keys----------------------------

pine_cone_key = os.getenv("PINE_CONE_API_KEY")
Gen_Ai_key = os.getenv("GEN_AI_API_KEY")


# client = genai.Client(api_key=Gen_Ai_key)



pc = Pinecone(api_key=pine_cone_key)

index = pc.Index("practice")



async def save_file_to_gridfs(file_bytes: bytes, filename: str, content_type: str):
    fs = AsyncIOMotorGridFSBucket(db)
    file_stream = BytesIO(file_bytes)
    file_id = await fs.upload_from_stream(filename, file_stream, metadata={"content_type": content_type})
    return str(file_id)



async def get_file_from_gridfs(file_id: str):
    fs = AsyncIOMotorGridFSBucket(db)
    oid = ObjectId(file_id)
    file_obj = await fs.open_download_stream(oid)
    return file_obj


def parse_uploaded_file(file_bytes: bytes, filename: str):
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, filename)

    with open(file_path, "wb") as f:
        f.write(file_bytes)

    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = "\n".join(
            page.extract_text() for page in reader.pages
            if page.extract_text()
        )

    elif ext == ".docx":
        import docx
        doc = docx.Document(file_path)
        text = "\n".join(
            para.text for para in doc.paragraphs
            if para.text.strip()
        )

    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return text

    # elements = partition(filename=file_path)
    # text = "\n".join(
    #         str(element)
    #         for element in elements
    #         if str(element).strip()
    #         )

    # return text
    # extracted_data = [str(element) for element in elements]

    # return extracted_data


# async def save_chunks_to_db(file_id: str, chunks: list[str]):
#     documents = [{"file_id": file_id, "chunk": chunk} for chunk in chunks]
#     await parsed_data_col.insert_many(documents)










@api_router.post("/Upload_file", tags=["File"])
async def upload_file(file: UploadFile = File(...)):
    try:
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/gemini-embedding-001",
            google_api_key=Gen_Ai_key
        )
        # embeddings = HuggingFaceEmbeddings(
        #     model_name="sentence-transformers/all-MiniLM-L6-v2"
        # )

        vector_store = PineconeVectorStore(
            index=index,
            embedding=embeddings
        )
        
        content = await file.read()

        file_id = await save_file_to_gridfs(content, file.filename, file.content_type)
        text = parse_uploaded_file(content, file.filename)
        
        print(text)
        documents = [
            Document(
                page_content=text,
                metadata={
                    "file_id": str(file_id),
                    "file_name": file.filename
                }
            )
        ]

        splitter = RecursiveCharacterTextSplitter(
                    chunk_size=500,
                    chunk_overlap=100
                    )

        chunks = splitter.split_documents(documents)

        vector_store.add_documents(chunks)

        return {
            "Message" : "Document Uploaded Successfully",
            "file_id": file_id,
            "parsed_content": text
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Something Went Wrong, Please check your network connection! {str(e)}")





# @api_router.get("/view_file", tags=["File"])
# async def get_file(id:str):
#     try:
#         oid = ObjectId(id)
#         file = await get_file_from_gridfs(oid)

#         file_info = await db.fs.files.find_one({"_id": oid})
#         if not file_info:
#             raise HTTPException(status_code=404, detail="File not found")
        
#         content_type = file_info.get("contentType", "application/octet-stream")
#         filename = file_info.get("filename", "file")
        
#         return StreamingResponse(
#             file,
#             media_type=content_type,
#             headers={"Content-Disposition": f"inline; filename={filename}"}
#         )
#     except Exception as e:
#         raise HTTPException(status_code=404, detail=str(e))
    













#----------------------------old code --------------------------------------------------------

# @api_router.post("/Upload_file", tags=["File"])
# async def upload_file(file: UploadFile = File(...)):
#     try:
#         content = await file.read()

#         file_id = await save_file_to_gridfs(content, file.filename, file.content_type)
#         parsed_data = parse_uploaded_file(content, file.filename)

#         chunk_docs = [{"file_id": file_id, "chunk": chunk} for chunk in parsed_data]
#         await parsed_data_col.insert_many(chunk_docs)

#         # await save_chunks_to_db(file_id, parsed_data)
#         await create_faiss_index(file_id)

#         return {
#             "message": "File uploaded successfully",
#             "file_id": file_id,
#             "parsed_content": parsed_data
#         }
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Something Went Wrong, Please check your network connection! {str(e)}")

